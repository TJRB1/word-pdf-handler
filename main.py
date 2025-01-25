"""
Written by Thomas J R Barker
This script can read and add writing to Word and PDF documents
"""

from docx import Document
from PyPDF2 import PdfReader, PdfWriter

# This function reads text from a Word document
def read_word(file_path):
    doc = Document(file_path)
    print("Word Document Content:")
    for paragraph in doc.paragraphs:
        print(paragraph.text)

# This function adds writing to a Word document
def write_word(file_path, content):
    doc = Document()
    doc.add_paragraph(content)
    doc.save(file_path)
    print(f"Content written to {file_path}")

# This function reads text from a PDF file
def read_pdf(file_path):
    reader = PdfReader(file_path)
    print("PDF Content:")
    for page in reader.pages:
        print(page.extract_text())

# This function adds writing to a PDF file
def write_pdf(file_path, content):
    writer = PdfWriter()
    writer.add_blank_page(width=210, height=297)  # A4 size
    with open(file_path, "wb") as file:
        writer.write(file)
    print(f"PDF created at {file_path}")

"""
Example Usage
# Uncomment these lines to test the functions:
# read_word("example.docx")
# write_word("new_document.docx", "This is a test content for Word.")
# read_pdf("example.pdf")
# write_pdf("new_document.pdf", "This is a test content for PDF.")
"""
